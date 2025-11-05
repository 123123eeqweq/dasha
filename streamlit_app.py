# -*- coding: utf-8 -*-
import streamlit as st
import pandas as pd
from docx import Document
import re
import io

st.set_page_config(
    page_title="✨ Обработка артикулов ✨",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Девчачий CSS
st.markdown("""
<style>
    .main {
        background: linear-gradient(135deg, #FFF0F5 0%, #FFE4E1 100%);
    }
    .stButton>button {
        background: linear-gradient(135deg, #FF1493 0%, #FF69B4 100%);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.5rem 2rem;
        font-weight: bold;
        font-size: 18px;
        box-shadow: 0 5px 15px rgba(255, 20, 147, 0.4);
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 20px rgba(255, 20, 147, 0.6);
    }
    h1 {
        color: #8B008B;
        text-align: center;
        text-shadow: 2px 2px 4px rgba(255, 105, 180, 0.3);
    }
    .stFileUploader {
        background: white;
        border-radius: 10px;
        padding: 10px;
    }
</style>
""", unsafe_allow_html=True)

def extract_articles_from_word(word_file):
    """Извлекает артикулы и украинские названия из Word файла"""
    articles_dict = {}
    
    try:
        doc = Document(word_file)
        article_pattern = re.compile(r'\b\d{11}\b')
        all_text = []
        
        # Собираем текст из параграфов
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text.append(text)
        
        # Собираем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        all_text.append(text)
        
        # Ищем артикулы и названия
        i = 0
        while i < len(all_text):
            line = all_text[i]
            article_match = article_pattern.search(line)
            
            if article_match:
                article = article_match.group()
                ukrainian_name = None
                
                # Ищем украинское название
                if i + 1 < len(all_text):
                    next_line = all_text[i + 1]
                    if not article_pattern.search(next_line) and next_line:
                        if re.search(r'[А-Яа-яЄєІіЇїҐґ]', next_line):
                            ukrainian_name = next_line
                
                if not ukrainian_name:
                    after_article = line[article_match.end():].strip()
                    if after_article and re.search(r'[А-Яа-яЄєІіЇїҐґ]', after_article):
                        ukrainian_name = after_article
                
                if not ukrainian_name and i > 0:
                    prev_line = all_text[i - 1]
                    if not article_pattern.search(prev_line) and re.search(r'[А-Яа-яЄєІіЇїҐґ]', prev_line):
                        ukrainian_name = prev_line
                
                if ukrainian_name:
                    if article not in articles_dict or len(ukrainian_name) > len(articles_dict[article]):
                        articles_dict[article] = ukrainian_name
            
            i += 1
        
        return articles_dict
        
    except Exception as e:
        raise Exception(f"Ошибка при чтении Word файла: {e}")

def merge_to_excel(excel_file, articles_dict):
    """Объединяет данные в Excel"""
    try:
        # Читаем Excel
        try:
            df = pd.read_excel(excel_file, engine='xlrd')
        except:
            df = pd.read_excel(excel_file, engine='openpyxl')
        
        # Удаляем пустые колонки (Unnamed)
        columns_to_drop = [col for col in df.columns if str(col).startswith('Unnamed')]
        if columns_to_drop:
            df = df.drop(columns=columns_to_drop)
        
        df = df.dropna(axis=1, how='all')
        
        # Ищем столбец с артикулами
        article_column = None
        for col in df.columns:
            col_lower = str(col).lower()
            if 'stok' in col_lower and 'kodu' in col_lower:
                article_column = col
                break
        
        if article_column is None:
            raise ValueError("Не найден столбец 'STOK KODU' в Excel файле")
        
        # Проверяем, есть ли уже колонка с украинскими названиями
        ukr_column_name = 'Українська назва'
        if ukr_column_name in df.columns:
            df[ukr_column_name] = ''
        else:
            df[ukr_column_name] = ''
        
        # Ищем совпадения
        matched_count = 0
        for idx, row in df.iterrows():
            article = str(row[article_column]).strip()
            
            if article in articles_dict:
                df.at[idx, ukr_column_name] = articles_dict[article]
                matched_count += 1
            else:
                article_clean = article.replace(' ', '').replace('-', '').replace('.', '')
                for art_key, name in articles_dict.items():
                    art_key_clean = art_key.replace(' ', '').replace('-', '').replace('.', '')
                    if article_clean == art_key_clean:
                        df.at[idx, ukr_column_name] = name
                        matched_count += 1
                        break
        
        # Финальная проверка - удаляем все Unnamed колонки
        columns_to_drop_final = [col for col in df.columns if str(col).startswith('Unnamed')]
        if columns_to_drop_final:
            df = df.drop(columns=columns_to_drop_final)
        
        # Сохраняем в BytesIO
        output = io.BytesIO()
        df.to_excel(output, index=False, engine='openpyxl')
        output.seek(0)
        
        return output, matched_count, len(df)
        
    except Exception as e:
        raise Exception(f"Ошибка при работе с Excel: {e}")

# Интерфейс
st.title("✨ Обработка артикулов ✨")
st.markdown("<p style='text-align: center; color: #DA70D6; font-style: italic;'>Word и Excel файлы</p>", unsafe_allow_html=True)

# Загрузка файлов
col1, col2 = st.columns(2)

with col1:
    st.subheader("📊 Excel файл")
    excel_file = st.file_uploader("Выберите Excel файл", type=['xls', 'xlsx'], key='excel')

with col2:
    st.subheader("📝 Word файлы")
    word_files = st.file_uploader("Выберите Word файлы (можно несколько)", type=['docx'], key='word', accept_multiple_files=True)

# Кнопка обработки
if st.button("✨ Обработать файлы ✨", type="primary", use_container_width=True):
    if excel_file is None:
        st.error("❌ Выберите Excel файл!")
    elif not word_files:
        st.error("❌ Добавьте хотя бы один Word файл!")
    else:
        with st.spinner("⏳ Обработка файлов..."):
            try:
                # Обрабатываем все Word файлы
                all_articles_dict = {}
                processed_count = 0
                
                for word_file in word_files:
                    try:
                        articles_dict = extract_articles_from_word(word_file)
                        for article, name in articles_dict.items():
                            if article not in all_articles_dict or len(name) > len(all_articles_dict[article]):
                                all_articles_dict[article] = name
                        processed_count += 1
                    except Exception as e:
                        st.warning(f"⚠️ Ошибка при обработке {word_file.name}: {str(e)}")
                
                if not all_articles_dict:
                    st.error("❌ Не удалось извлечь артикулы из Word файлов!")
                else:
                    # Объединяем с Excel
                    output, matched_count, total = merge_to_excel(excel_file, all_articles_dict)
                    
                    # Показываем результат
                    st.success(f"✅ Обработка завершена! Найдено совпадений: {matched_count} из {total}")
                    
                    # Скачивание
                    st.download_button(
                        label="💾 Скачать результат",
                        data=output,
                        file_name="excel_with_ukrainian_names.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                    
            except Exception as e:
                st.error(f"❌ Ошибка: {str(e)}")

