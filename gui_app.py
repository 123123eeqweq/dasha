# -*- coding: utf-8 -*-
import sys
import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
from docx import Document
import pandas as pd
import re

# Устанавливаем UTF-8 для Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

class ExcelWordProcessor:
    def __init__(self, root):
        self.root = root
        self.root.title("✨ Обработка файлов ✨")
        self.root.geometry("800x700")
        
        # Девчачья цветовая палитра
        self.colors = {
            'bg_main': '#FFF0F5',  # Лавандовый бланш
            'bg_frame': '#FFE4E1',  # Мятный крем
            'accent_pink': '#FF69B4',  # Горячий розовый
            'accent_purple': '#DA70D6',  # Орхидея
            'accent_light': '#FFB6C1',  # Светло-розовый
            'text_dark': '#8B008B',  # Темно-фиолетовый
            'button_pink': '#FF1493',  # Глубокий розовый
            'button_purple': '#BA55D3',  # Средняя орхидея
            'button_light': '#FFB6C1',  # Светло-розовый
            'success': '#FF69B4',  # Розовый для успеха
        }
        
        # Устанавливаем фон окна
        self.root.configure(bg=self.colors['bg_main'])
        
        # Переменные для путей к файлам
        self.excel_path = tk.StringVar()
        self.word_paths = []  # Список путей к Word файлам
        self.output_path = tk.StringVar()
        
        # Создаем интерфейс
        self.create_widgets()
        
    def create_widgets(self):
        # Заголовок с эмодзи
        header_frame = tk.Frame(self.root, bg=self.colors['bg_main'])
        header_frame.pack(pady=15)
        
        title_label = tk.Label(
            header_frame, 
            text="✨ Обработка артикулов ✨",
            font=("Segoe UI", 20, "bold"),
            fg=self.colors['text_dark'],
            bg=self.colors['bg_main'],
            pady=5
        )
        title_label.pack()
        
        subtitle_label = tk.Label(
            header_frame,
            text="Word и Excel файлы",
            font=("Segoe UI", 11, "italic"),
            fg=self.colors['accent_purple'],
            bg=self.colors['bg_main']
        )
        subtitle_label.pack()
        
        # Фрейм для Excel файла
        excel_frame = tk.LabelFrame(
            self.root, 
            text="📊 1. Excel файл",
            padx=15,
            pady=12,
            font=("Segoe UI", 10, "bold"),
            fg=self.colors['text_dark'],
            bg=self.colors['bg_frame'],
            relief=tk.RAISED,
            bd=2
        )
        excel_frame.pack(fill=tk.X, padx=25, pady=8)
        
        excel_entry = tk.Entry(
            excel_frame, 
            textvariable=self.excel_path, 
            width=45,
            font=("Segoe UI", 9),
            relief=tk.SUNKEN,
            bd=2
        )
        excel_entry.pack(side=tk.LEFT, padx=8, pady=5)
        
        excel_btn = tk.Button(
            excel_frame, 
            text="📁 Выбрать",
            command=lambda: self.select_file("excel", [("Excel files", "*.xls *.xlsx")]),
            bg=self.colors['button_pink'],
            fg="white",
            font=("Segoe UI", 9, "bold"),
            relief=tk.RAISED,
            bd=3,
            padx=12,
            pady=5,
            cursor="hand2",
            activebackground=self.colors['accent_pink'],
            activeforeground="white"
        )
        excel_btn.pack(side=tk.LEFT, padx=5)
        
        # Фрейм для Word файлов (можно несколько)
        word_frame = tk.LabelFrame(
            self.root, 
            text="📝 2. Word файлы (можно несколько)",
            padx=15,
            pady=12,
            font=("Segoe UI", 10, "bold"),
            fg=self.colors['text_dark'],
            bg=self.colors['bg_frame'],
            relief=tk.RAISED,
            bd=2
        )
        word_frame.pack(fill=tk.BOTH, expand=True, padx=25, pady=8)
        
        # Кнопки управления
        word_buttons_frame = tk.Frame(word_frame, bg=self.colors['bg_frame'])
        word_buttons_frame.pack(fill=tk.X, pady=8)
        
        add_word_btn = tk.Button(
            word_buttons_frame, 
            text="➕ Добавить файлы",
            command=lambda: self.add_word_file(),
            bg=self.colors['button_purple'],
            fg="white",
            font=("Segoe UI", 9, "bold"),
            relief=tk.RAISED,
            bd=3,
            padx=12,
            pady=6,
            cursor="hand2",
            activebackground=self.colors['accent_purple'],
            activeforeground="white"
        )
        add_word_btn.pack(side=tk.LEFT, padx=5)
        
        remove_word_btn = tk.Button(
            word_buttons_frame,
            text="➖ Удалить",
            command=self.remove_selected_word_file,
            bg=self.colors['button_light'],
            fg=self.colors['text_dark'],
            font=("Segoe UI", 9, "bold"),
            relief=tk.RAISED,
            bd=3,
            padx=12,
            pady=6,
            cursor="hand2",
            activebackground=self.colors['accent_light'],
            activeforeground=self.colors['text_dark']
        )
        remove_word_btn.pack(side=tk.LEFT, padx=5)
        
        # Список файлов с прокруткой
        list_frame = tk.Frame(word_frame, bg=self.colors['bg_frame'])
        list_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        scrollbar_word = tk.Scrollbar(list_frame, bg=self.colors['accent_light'])
        scrollbar_word.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.word_listbox = tk.Listbox(
            list_frame, 
            yscrollcommand=scrollbar_word.set, 
            height=4,
            font=("Segoe UI", 9),
            bg="white",
            fg=self.colors['text_dark'],
            selectbackground=self.colors['accent_light'],
            selectforeground=self.colors['text_dark'],
            relief=tk.SUNKEN,
            bd=2
        )
        self.word_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_word.config(command=self.word_listbox.yview)
        
        # Фрейм для выходного файла
        output_frame = tk.LabelFrame(
            self.root, 
            text="💾 3. Сохранить результат",
            padx=15,
            pady=12,
            font=("Segoe UI", 10, "bold"),
            fg=self.colors['text_dark'],
            bg=self.colors['bg_frame'],
            relief=tk.RAISED,
            bd=2
        )
        output_frame.pack(fill=tk.X, padx=25, pady=8)
        
        output_entry = tk.Entry(
            output_frame, 
            textvariable=self.output_path, 
            width=45,
            font=("Segoe UI", 9),
            relief=tk.SUNKEN,
            bd=2
        )
        output_entry.pack(side=tk.LEFT, padx=8, pady=5)
        
        output_btn = tk.Button(
            output_frame, 
            text="📂 Выбрать",
            command=self.select_output_path,
            bg=self.colors['button_purple'],
            fg="white",
            font=("Segoe UI", 9, "bold"),
            relief=tk.RAISED,
            bd=3,
            padx=12,
            pady=5,
            cursor="hand2",
            activebackground=self.colors['accent_purple'],
            activeforeground="white"
        )
        output_btn.pack(side=tk.LEFT, padx=5)
        
        # Кнопка обработки - большая и красивая
        button_frame = tk.Frame(self.root, bg=self.colors['bg_main'])
        button_frame.pack(pady=15)
        
        process_btn = tk.Button(
            button_frame,
            text="✨ Обработать файлы ✨",
            command=self.process_files,
            bg=self.colors['button_pink'],
            fg="white",
            font=("Segoe UI", 13, "bold"),
            relief=tk.RAISED,
            bd=4,
            padx=30,
            pady=12,
            cursor="hand2",
            activebackground=self.colors['accent_pink'],
            activeforeground="white"
        )
        process_btn.pack()
        
        # Прогресс бар с красивым стилем
        progress_frame = tk.Frame(self.root, bg=self.colors['bg_main'])
        progress_frame.pack(pady=10)
        
        self.progress = ttk.Progressbar(
            progress_frame, 
            mode='indeterminate',
            length=650,
            style="TProgressbar"
        )
        self.progress.pack()
        
        # Настраиваем стиль прогресс-бара
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("TProgressbar",
                        background=self.colors['accent_pink'],
                        troughcolor=self.colors['bg_frame'],
                        borderwidth=0,
                        lightcolor=self.colors['accent_pink'],
                        darkcolor=self.colors['accent_pink'])
        
        # Текстовое поле для логов
        log_frame = tk.LabelFrame(
            self.root, 
            text="📋 Лог обработки",
            padx=15,
            pady=12,
            font=("Segoe UI", 10, "bold"),
            fg=self.colors['text_dark'],
            bg=self.colors['bg_frame'],
            relief=tk.RAISED,
            bd=2
        )
        log_frame.pack(fill=tk.BOTH, expand=True, padx=25, pady=8)
        
        self.log_text = tk.Text(
            log_frame, 
            height=8, 
            wrap=tk.WORD,
            font=("Segoe UI", 9),
            bg="white",
            fg=self.colors['text_dark'],
            relief=tk.SUNKEN,
            bd=2,
            padx=5,
            pady=5
        )
        scrollbar = tk.Scrollbar(
            log_frame, 
            orient=tk.VERTICAL, 
            command=self.log_text.yview,
            bg=self.colors['accent_light']
        )
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
    def log(self, message):
        """Добавляет сообщение в лог"""
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update()
        
    def select_file(self, file_type, filetypes):
        """Выбор файла"""
        filename = filedialog.askopenfilename(filetypes=filetypes)
        if filename:
            if file_type == "excel":
                self.excel_path.set(filename)
                self.log(f"Выбран Excel файл: {os.path.basename(filename)}")
    
    def add_word_file(self):
        """Добавление Word файла в список"""
        filenames = filedialog.askopenfilenames(
            title="Выберите Word файлы (можно несколько)",
            filetypes=[("Word files", "*.docx")]
        )
        for filename in filenames:
            if filename not in self.word_paths:
                self.word_paths.append(filename)
                self.word_listbox.insert(tk.END, os.path.basename(filename))
                self.log(f"Добавлен Word файл: {os.path.basename(filename)}")
    
    def remove_selected_word_file(self):
        """Удаление выбранного Word файла из списка"""
        selected = self.word_listbox.curselection()
        if selected:
            index = selected[0]
            removed_file = self.word_paths.pop(index)
            self.word_listbox.delete(index)
            self.log(f"Удален Word файл: {os.path.basename(removed_file)}")
        else:
            messagebox.showwarning("Предупреждение", "Выберите файл для удаления из списка")
    
    def select_output_path(self):
        """Выбор пути для сохранения"""
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("Excel files", "*.xls")]
        )
        if filename:
            self.output_path.set(filename)
            self.log(f"Результат будет сохранен: {os.path.basename(filename)}")
    
    def extract_articles_from_word(self, word_path):
        """Извлекает артикулы и украинские названия из одного Word файла"""
        articles_dict = {}
        
        try:
            doc = Document(word_path)
            article_pattern = re.compile(r'\b\d{11}\b')
            all_text = []
            
            # Собираем текст из параграфов
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    all_text.append(text)
            
            # Собираем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            all_text.append(text)
            
            # Ищем артикулы и названия
            i = 0
            while i < len(all_text):
                line = all_text[i]
                article_match = article_pattern.search(line)
                
                if article_match:
                    article = article_match.group()
                    ukrainian_name = None
                    
                    # Ищем украинское название
                    if i + 1 < len(all_text):
                        next_line = all_text[i + 1]
                        if not article_pattern.search(next_line) and next_line:
                            if re.search(r'[А-Яа-яЄєІіЇїҐґ]', next_line):
                                ukrainian_name = next_line
                    
                    if not ukrainian_name:
                        after_article = line[article_match.end():].strip()
                        if after_article and re.search(r'[А-Яа-яЄєІіЇїҐґ]', after_article):
                            ukrainian_name = after_article
                    
                    if not ukrainian_name and i > 0:
                        prev_line = all_text[i - 1]
                        if not article_pattern.search(prev_line) and re.search(r'[А-Яа-яЄєІіЇїҐґ]', prev_line):
                            ukrainian_name = prev_line
                    
                    if ukrainian_name:
                        # Если артикул уже есть, берем более длинное название (более полное)
                        if article not in articles_dict or len(ukrainian_name) > len(articles_dict[article]):
                            articles_dict[article] = ukrainian_name
                
                i += 1
            
            return articles_dict
            
        except Exception as e:
            raise Exception(f"Ошибка при чтении {os.path.basename(word_path)}: {e}")
    
    def extract_articles_from_all_words(self, word_paths):
        """Извлекает артикулы из всех Word файлов и объединяет их"""
        self.log(f"Извлечение артикулов из {len(word_paths)} Word файлов...")
        all_articles_dict = {}
        
        for i, word_path in enumerate(word_paths, 1):
            filename = os.path.basename(word_path)
            self.log(f"Обработка файла {i}/{len(word_paths)}: {filename}...")
            
            try:
                articles_dict = self.extract_articles_from_word(word_path)
                found_count = len(articles_dict)
                
                # Объединяем словари
                for article, name in articles_dict.items():
                    # Если артикул уже есть, берем более длинное название
                    if article not in all_articles_dict or len(name) > len(all_articles_dict[article]):
                        all_articles_dict[article] = name
                
                self.log(f"  ✓ Найдено в {filename}: {found_count} артикулов")
                
            except Exception as e:
                self.log(f"  ✗ Ошибка в {filename}: {e}")
                # Продолжаем обработку других файлов
        
        self.log(f"Всего уникальных артикулов из всех файлов: {len(all_articles_dict)}")
        return all_articles_dict
    
    def merge_to_excel(self, excel_path, articles_dict, output_path):
        """Объединяет данные в Excel"""
        self.log("Чтение Excel файла...")
        
        try:
            # Читаем Excel
            try:
                df = pd.read_excel(excel_path, engine='xlrd')
            except:
                df = pd.read_excel(excel_path, engine='openpyxl')
            
            self.log(f"Загружено строк из Excel: {len(df)}")
            self.log(f"Колонки до очистки: {list(df.columns)}")
            
            # Удаляем пустые колонки (Unnamed) - более надежный способ
            columns_to_drop = [col for col in df.columns if str(col).startswith('Unnamed')]
            if columns_to_drop:
                self.log(f"Удаляем пустые колонки: {columns_to_drop}")
                df = df.drop(columns=columns_to_drop)
            
            # Также удаляем полностью пустые колонки
            df = df.dropna(axis=1, how='all')
            
            self.log(f"Колонки после очистки: {list(df.columns)}")
            
            # Ищем столбец с артикулами
            article_column = None
            for col in df.columns:
                col_lower = str(col).lower()
                if 'stok' in col_lower and 'kodu' in col_lower:
                    article_column = col
                    break
            
            if article_column is None:
                raise ValueError("Не найден столбец 'STOK KODU' в Excel файле")
            
            self.log(f"Найден столбец с артикулами: '{article_column}'")
            
            # Проверяем, есть ли уже колонка с украинскими названиями
            ukr_column_name = 'Українська назва'
            if ukr_column_name in df.columns:
                self.log("Колонка 'Українська назва' уже существует, обновляем её...")
                # Очищаем существующую колонку
                df[ukr_column_name] = ''
            else:
                # Добавляем новую колонку
                df[ukr_column_name] = ''
                self.log("Добавлена новая колонка 'Українська назва'")
            
            # Ищем совпадения
            matched_count = 0
            for idx, row in df.iterrows():
                article = str(row[article_column]).strip()
                
                if article in articles_dict:
                    df.at[idx, ukr_column_name] = articles_dict[article]
                    matched_count += 1
                else:
                    article_clean = article.replace(' ', '').replace('-', '').replace('.', '')
                    for art_key, name in articles_dict.items():
                        art_key_clean = art_key.replace(' ', '').replace('-', '').replace('.', '')
                        if article_clean == art_key_clean:
                            df.at[idx, ukr_column_name] = name
                            matched_count += 1
                            break
            
            self.log(f"Найдено совпадений: {matched_count} из {len(df)}")
            
            # Финальная проверка - удаляем все Unnamed колонки перед сохранением
            columns_to_drop_final = [col for col in df.columns if str(col).startswith('Unnamed')]
            if columns_to_drop_final:
                self.log(f"Финальная очистка: удаляем {columns_to_drop_final}")
                df = df.drop(columns=columns_to_drop_final)
            
            self.log(f"Финальные колонки перед сохранением: {list(df.columns)}")
            
            # Сохраняем
            self.log("Сохранение результата...")
            df.to_excel(output_path, index=False, engine='openpyxl')
            self.log(f"✓ Файл сохранен: {os.path.basename(output_path)}")
            
            return matched_count, len(df)
            
        except Exception as e:
            self.log(f"Ошибка при работе с Excel: {e}")
            raise
    
    def process_files(self):
        """Основная функция обработки"""
        # Проверяем файлы
        if not self.excel_path.get():
            messagebox.showerror("Ошибка", "Выберите Excel файл!")
            return
        
        if not self.word_paths:
            messagebox.showerror("Ошибка", "Добавьте хотя бы один Word файл!")
            return
        
        if not self.output_path.get():
            messagebox.showerror("Ошибка", "Укажите путь для сохранения результата!")
            return
        
        # Очищаем лог
        self.log_text.delete(1.0, tk.END)
        self.log("Начало обработки...")
        self.log(f"Excel файл: {os.path.basename(self.excel_path.get())}")
        self.log(f"Word файлов: {len(self.word_paths)}")
        self.log("=" * 50)
        
        # Запускаем прогресс бар
        self.progress.start()
        
        # Запускаем обработку в отдельном потоке
        thread = threading.Thread(target=self.process_in_thread)
        thread.daemon = True
        thread.start()
    
    def process_in_thread(self):
        """Обработка в отдельном потоке"""
        try:
            # Извлекаем артикулы из всех Word файлов
            articles_dict = self.extract_articles_from_all_words(self.word_paths)
            
            # Объединяем с Excel
            matched, total = self.merge_to_excel(
                self.excel_path.get(),
                articles_dict,
                self.output_path.get()
            )
            
            # Останавливаем прогресс
            self.progress.stop()
            
            # Показываем результат
            self.log("=" * 50)
            self.log("✓ Обработка завершена успешно!")
            self.log(f"Результат: {matched} из {total} артикулов получили украинские названия")
            
            messagebox.showinfo(
                "Успех",
                f"Обработка завершена!\n\n"
                f"Найдено совпадений: {matched} из {total}\n"
                f"Файл сохранен: {os.path.basename(self.output_path.get())}"
            )
            
        except Exception as e:
            self.progress.stop()
            self.log(f"✗ Ошибка: {e}")
            messagebox.showerror("Ошибка", f"Произошла ошибка:\n{e}")

def main():
    root = tk.Tk()
    app = ExcelWordProcessor(root)
    root.mainloop()

if __name__ == "__main__":
    main()

