from docx import Document

def read_word_file(file_path):
    """Читает Word файл и выводит его содержимое"""
    try:
        # Открываем документ
        doc = Document(file_path)
        
        print(f"Файл: {file_path}")
        print("-" * 50)
        
        # Читаем все параграфы
        print("\nТекст документа:")
        print("=" * 50)
        for i, paragraph in enumerate(doc.paragraphs, 1):
            if paragraph.text.strip():  # Пропускаем пустые параграфы
                print(f"{i}. {paragraph.text}")
        
        # Читаем таблицы, если есть
        if len(doc.tables) > 0:
            print("\n" + "=" * 50)
            print(f"Таблицы ({len(doc.tables)} шт.):")
            print("=" * 50)
            
            for table_idx, table in enumerate(doc.tables, 1):
                print(f"\nТаблица {table_idx}:")
                print("-" * 50)
                
                for row_idx, row in enumerate(table.rows, 1):
                    row_data = [cell.text.strip() for cell in row.cells]
                    print(f"Строка {row_idx}: {' | '.join(row_data)}")
        
        # Информация о документе
        print("\n" + "=" * 50)
        print("Информация:")
        print("=" * 50)
        print(f"Всего параграфов: {len(doc.paragraphs)}")
        print(f"Всего таблиц: {len(doc.tables)}")
        
        # Проверяем стили
        if doc.styles:
            print(f"Доступных стилей: {len(doc.styles)}")
        
    except Exception as e:
        print(f"Ошибка при чтении файла: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    # Читаем файл word.docx
    read_word_file("word.docx")

