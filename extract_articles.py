# -*- coding: utf-8 -*-
import sys
import re
from docx import Document

# Устанавливаем UTF-8 для вывода в Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

def extract_articles_and_names(file_path):
    """Извлекает артикулы и украинские названия из Word документа"""
    try:
        doc = Document(file_path)
        
        # Регулярное выражение для поиска артикулов (последовательность из 11 цифр)
        article_pattern = re.compile(r'\b\d{11}\b')
        
        results = []
        all_text = []
        
        # Собираем весь текст из параграфов
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text.append(text)
        
        # Собираем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        all_text.append(text)
        
        # Ищем артикулы и связанные названия
        i = 0
        while i < len(all_text):
            line = all_text[i]
            
            # Ищем артикул в строке
            article_match = article_pattern.search(line)
            
            if article_match:
                article = article_match.group()
                
                # Ищем украинское название - обычно следующая строка или в той же строке после артикула
                ukrainian_name = None
                
                # Вариант 1: название в следующей строке
                if i + 1 < len(all_text):
                    next_line = all_text[i + 1]
                    # Если следующая строка не артикул и не пустая, возможно это название
                    if not article_pattern.search(next_line) and next_line:
                        # Проверяем, содержит ли строка кириллицу (украинский текст)
                        if re.search(r'[А-Яа-яЄєІіЇїҐґ]', next_line):
                            ukrainian_name = next_line
                
                # Вариант 2: название в той же строке после артикула
                if not ukrainian_name:
                    after_article = line[article_match.end():].strip()
                    if after_article and re.search(r'[А-Яа-яЄєІіЇїҐґ]', after_article):
                        ukrainian_name = after_article
                
                # Вариант 3: название в предыдущей строке (если артикул на отдельной строке)
                if not ukrainian_name and i > 0:
                    prev_line = all_text[i - 1]
                    if not article_pattern.search(prev_line) and re.search(r'[А-Яа-яЄєІіЇїҐґ]', prev_line):
                        ukrainian_name = prev_line
                
                if ukrainian_name:
                    results.append({
                        'article': article,
                        'ukrainian_name': ukrainian_name
                    })
                    print(f"Артикул: {article}")
                    print(f"Украинское название: {ukrainian_name}")
                    print("-" * 50)
            
            i += 1
        
        print(f"\nВсего найдено: {len(results)} артикулов")
        return results
        
    except Exception as e:
        print(f"Ошибка: {e}")
        import traceback
        traceback.print_exc()
        return []

if __name__ == "__main__":
    results = extract_articles_and_names("word.docx")
    
    # Сохраняем результаты в файл
    if results:
        with open('articles_output.txt', 'w', encoding='utf-8') as f:
            for item in results:
                f.write(f"{item['article']}\t{item['ukrainian_name']}\n")
        print("\nРезультаты сохранены в articles_output.txt")

