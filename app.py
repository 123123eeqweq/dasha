# -*- coding: utf-8 -*-
from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from werkzeug.utils import secure_filename
import os
import pandas as pd
from docx import Document
import re
from io import BytesIO
import tempfile

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Создаем папку для загрузок
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'xls', 'xlsx', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_articles_from_word(word_path):
    """Извлекает артикулы и украинские названия из Word файла"""
    articles_dict = {}
    
    try:
        doc = Document(word_path)
        article_pattern = re.compile(r'\b\d{11}\b')
        all_text = []
        
        # Собираем текст из параграфов
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text.append(text)
        
        # Собираем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        all_text.append(text)
        
        # Ищем артикулы и названия
        i = 0
        while i < len(all_text):
            line = all_text[i]
            article_match = article_pattern.search(line)
            
            if article_match:
                article = article_match.group()
                ukrainian_name = None
                
                # Ищем украинское название
                if i + 1 < len(all_text):
                    next_line = all_text[i + 1]
                    if not article_pattern.search(next_line) and next_line:
                        if re.search(r'[А-Яа-яЄєІіЇїҐґ]', next_line):
                            ukrainian_name = next_line
                
                if not ukrainian_name:
                    after_article = line[article_match.end():].strip()
                    if after_article and re.search(r'[А-Яа-яЄєІіЇїҐґ]', after_article):
                        ukrainian_name = after_article
                
                if not ukrainian_name and i > 0:
                    prev_line = all_text[i - 1]
                    if not article_pattern.search(prev_line) and re.search(r'[А-Яа-яЄєІіЇїҐґ]', prev_line):
                        ukrainian_name = prev_line
                
                if ukrainian_name:
                    if article not in articles_dict or len(ukrainian_name) > len(articles_dict[article]):
                        articles_dict[article] = ukrainian_name
            
            i += 1
        
        return articles_dict
        
    except Exception as e:
        raise Exception(f"Ошибка при чтении Word файла: {e}")

def merge_to_excel(excel_path, articles_dict):
    """Объединяет данные в Excel"""
    try:
        # Читаем Excel
        try:
            df = pd.read_excel(excel_path, engine='xlrd')
        except:
            df = pd.read_excel(excel_path, engine='openpyxl')
        
        # Удаляем пустые колонки (Unnamed)
        columns_to_drop = [col for col in df.columns if str(col).startswith('Unnamed')]
        if columns_to_drop:
            df = df.drop(columns=columns_to_drop)
        
        df = df.dropna(axis=1, how='all')
        
        # Ищем столбец с артикулами
        article_column = None
        for col in df.columns:
            col_lower = str(col).lower()
            if 'stok' in col_lower and 'kodu' in col_lower:
                article_column = col
                break
        
        if article_column is None:
            raise ValueError("Не найден столбец 'STOK KODU' в Excel файле")
        
        # Проверяем, есть ли уже колонка с украинскими названиями
        ukr_column_name = 'Українська назва'
        if ukr_column_name in df.columns:
            df[ukr_column_name] = ''
        else:
            df[ukr_column_name] = ''
        
        # Ищем совпадения
        matched_count = 0
        for idx, row in df.iterrows():
            article = str(row[article_column]).strip()
            
            if article in articles_dict:
                df.at[idx, ukr_column_name] = articles_dict[article]
                matched_count += 1
            else:
                article_clean = article.replace(' ', '').replace('-', '').replace('.', '')
                for art_key, name in articles_dict.items():
                    art_key_clean = art_key.replace(' ', '').replace('-', '').replace('.', '')
                    if article_clean == art_key_clean:
                        df.at[idx, ukr_column_name] = name
                        matched_count += 1
                        break
        
        # Финальная проверка - удаляем все Unnamed колонки
        columns_to_drop_final = [col for col in df.columns if str(col).startswith('Unnamed')]
        if columns_to_drop_final:
            df = df.drop(columns=columns_to_drop_final)
        
        # Сохраняем во временный файл
        output = BytesIO()
        df.to_excel(output, index=False, engine='openpyxl')
        output.seek(0)
        
        return output, matched_count, len(df)
        
    except Exception as e:
        raise Exception(f"Ошибка при работе с Excel: {e}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_files():
    try:
        # Проверяем файлы
        if 'excel_file' not in request.files:
            flash('Выберите Excel файл!', 'error')
            return redirect(url_for('index'))
        
        if 'word_files' not in request.files:
            flash('Добавьте хотя бы один Word файл!', 'error')
            return redirect(url_for('index'))
        
        excel_file = request.files['excel_file']
        word_files = request.files.getlist('word_files')
        
        if excel_file.filename == '':
            flash('Выберите Excel файл!', 'error')
            return redirect(url_for('index'))
        
        if not word_files or all(f.filename == '' for f in word_files):
            flash('Добавьте хотя бы один Word файл!', 'error')
            return redirect(url_for('index'))
        
        # Сохраняем файлы
        excel_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(excel_file.filename))
        excel_file.save(excel_path)
        
        # Обрабатываем все Word файлы
        all_articles_dict = {}
        processed_files = []
        
        for word_file in word_files:
            if word_file.filename:
                word_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(word_file.filename))
                word_file.save(word_path)
                
                try:
                    articles_dict = extract_articles_from_word(word_path)
                    for article, name in articles_dict.items():
                        if article not in all_articles_dict or len(name) > len(all_articles_dict[article]):
                            all_articles_dict[article] = name
                    processed_files.append(word_file.filename)
                except Exception as e:
                    flash(f'Ошибка при обработке {word_file.filename}: {str(e)}', 'warning')
                finally:
                    # Удаляем временный файл
                    if os.path.exists(word_path):
                        os.remove(word_path)
        
        if not all_articles_dict:
            flash('Не удалось извлечь артикулы из Word файлов!', 'error')
            if os.path.exists(excel_path):
                os.remove(excel_path)
            return redirect(url_for('index'))
        
        # Объединяем с Excel
        output, matched_count, total = merge_to_excel(excel_path, all_articles_dict)
        
        # Удаляем временный Excel файл
        if os.path.exists(excel_path):
            os.remove(excel_path)
        
        # Отправляем результат
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='excel_with_ukrainian_names.xlsx'
        )
        
    except Exception as e:
        flash(f'Ошибка: {str(e)}', 'error')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)

